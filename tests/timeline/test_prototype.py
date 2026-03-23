from pathlib import Path
import json
from datetime import datetime
from uuid import uuid4

import pytest

from ansimon_ai.llm.mock import MockLLMClient
from ansimon_ai.ocr.types import OCRResult, OCRSegment
from ansimon_ai.stt.mock import MockSTT
from ansimon_ai.timeline import (
    IncidentLogFormInput,
    TimelinePrototypeAiInput,
    TimelinePrototypeEvidenceInput,
    build_timeline_event_evidences,
    build_timeline_prototype,
)

TEST_TMP_DIR = Path("data/_timeline_test_tmp")

class SummaryLLMClient:
    def generate(self, messages: list[dict]) -> str:
        return json.dumps(
            {
                "evidence_metadata": {
                    "value": {
                        "evidence_type": "text",
                        "source": "unknown",
                        "sources": ["unknown"],
                        "created_at": "unknown",
                    },
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "parties": {
                    "value": {
                        "actor": "unknown",
                        "target": "unknown",
                        "relationship": "unknown",
                    },
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "period": {
                    "value": "unknown",
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "frequency": {
                    "value": "unknown",
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "channel": {
                    "value": ["unknown"],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "locations": {
                    "value": ["unknown"],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "action_types": {
                    "value": [],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "refusal_signal": {
                    "value": "unknown",
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "threat_indicators": {
                    "value": [],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "impact_on_victim": {
                    "value": [],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "report_or_record": {
                    "value": "unknown",
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "timeline_summary": {
                    "value": {
                        "title": "반복 연락 위협",
                        "description": "상대방이 반복적으로 연락하며 위협성 발언을 한 정황입니다.",
                    },
                    "confidence": "medium",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
            },
            ensure_ascii=False,
        )

def _write_test_file(name: str, content: bytes) -> Path:
    TEST_TMP_DIR.mkdir(parents=True, exist_ok=True)
    path = TEST_TMP_DIR / name
    path.write_bytes(content)
    return path

def test_build_timeline_prototype_completes_incident_log_and_skips_victim():
    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="INCIDENT_LOG",
                incident_log_form=IncidentLogFormInput(
                    title="incident title",
                    date="2026-03-19",
                    time="21:10",
                    place="home",
                    situation="the actor repeatedly knocked on the door",
                ),
            ),
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="VICTIM",
                file_format="IMAGE",
            ),
        ],
    )

    result = build_timeline_prototype(payload, llm_client=MockLLMClient())

    assert result.model_version == "prototype-v1"
    assert len(result.evidence_results) == 2

    completed = next(item for item in result.evidence_results if item.type == "INCIDENT_LOG")
    skipped = next(item for item in result.evidence_results if item.type == "VICTIM")

    assert completed.status == "completed"
    assert completed.source_type == "form"
    assert completed.title == "incident title"
    assert completed.structured_data is not None
    assert completed.timestamp is not None
    assert completed.timestamp.isoformat() == "2026-03-19T21:10:00"

    assert skipped.status == "skipped"
    assert skipped.error_code == "UNSUPPORTED_EVIDENCE_TYPE"

    assert len(result.items) == 1
    assert result.items[0].date == "2026-03-19"
    assert len(result.items[0].events) == 1
    timeline_evidence = result.items[0].events[0].evidences[0]
    assert timeline_evidence.title == "incident title"
    assert timeline_evidence.referenced_evidence_count == 1
    assert timeline_evidence.referenced_evidence_ids == [completed.evidence_id]

def test_build_timeline_prototype_prefers_llm_summary_for_title_and_description():
    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="REPORT_RECORD",
                file_format="TXT",
                extracted_text="2026-03-19 repeated threatening messages were documented.",
            ),
        ],
    )

    result = build_timeline_prototype(payload, llm_client=SummaryLLMClient())

    evidence_result = result.evidence_results[0]
    timeline_evidence = result.items[0].events[0].evidences[0]

    assert evidence_result.title == "반복 연락 위협"
    assert evidence_result.description == "상대방이 반복적으로 연락하며 위협성 발언을 한 정황입니다."
    assert timeline_evidence.title == evidence_result.title
    assert timeline_evidence.description == evidence_result.description

def test_build_timeline_prototype_uses_extracted_text_for_report_record():
    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="REPORT_RECORD",
                file_format="TXT",
                extracted_text="2026-03-18 consultation record\nrepeated contact was documented.",
            ),
        ],
    )

    result = build_timeline_prototype(payload, llm_client=MockLLMClient())

    assert len(result.evidence_results) == 1
    evidence_result = result.evidence_results[0]

    assert evidence_result.status == "completed"
    assert evidence_result.source_type == "document"
    assert evidence_result.normalized_text is not None
    assert result.items[0].date == "2026-03-18"

def test_build_timeline_prototype_skips_hwp_report_record_without_extracted_text():
    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="REPORT_RECORD",
                file_format="HWP",
                file_name="sample.hwp",
                file_bytes=b"fake-hwp",
            ),
        ],
    )

    result = build_timeline_prototype(payload, llm_client=MockLLMClient())

    evidence_result = result.evidence_results[0]
    assert evidence_result.status == "skipped"
    assert evidence_result.error_code == "UNSUPPORTED_FILE_FORMAT"
    assert evidence_result.error_message is not None
    assert "not supported yet" in evidence_result.error_message
    assert result.items == []

def test_build_timeline_prototype_accepts_hwp_report_record_with_extracted_text():
    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="REPORT_RECORD",
                file_format="HWP",
                extracted_text="2026-03-16 document extraction from hwp content",
                file_name="sample.hwp",
            ),
        ],
    )

    result = build_timeline_prototype(payload, llm_client=MockLLMClient())

    evidence_result = result.evidence_results[0]
    assert evidence_result.status == "completed"
    assert evidence_result.source_type == "document"
    assert result.items[0].date == "2026-03-16"

def test_build_timeline_prototype_processes_incident_log_document_upload():
    txt_path = _write_test_file(
        f"{uuid4()}-incident-log.txt",
        "2026-03-14\nincident log document upload\nactor appeared near workplace".encode("utf-8"),
    )

    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="INCIDENT_LOG",
                file_format="TXT",
                file_name="incident-log.txt",
                file_bytes=txt_path.read_bytes(),
            ),
        ],
    )

    result = build_timeline_prototype(payload, llm_client=MockLLMClient())

    evidence_result = result.evidence_results[0]
    timeline_evidence = result.items[0].events[0].evidences[0]
    assert evidence_result.status == "completed"
    assert evidence_result.source_type == "document"
    assert evidence_result.title
    assert evidence_result.title == timeline_evidence.title
    assert result.items[0].date == "2026-03-14"
    assert timeline_evidence.referenced_evidence_ids == [evidence_result.evidence_id]

def test_build_timeline_event_evidences_accumulates_referenced_ids():
    first_id = uuid4()
    second_id = uuid4()

    grouped = build_timeline_event_evidences(
        [
            {
                "evidence_id": first_id,
                "evidence_type": "MESSAGE",
                "message_group_key": "thread-1",
                "title": "message evidence",
                "description": "first",
                "tags": ["repeat"],
            },
            {
                "evidence_id": second_id,
                "evidence_type": "MESSAGE",
                "message_group_key": "thread-1",
                "title": "message evidence",
                "description": "second",
                "tags": ["threat"],
            },
        ]
    )

    assert len(grouped) == 1
    assert grouped[0]["referenced_evidence_count"] == 2
    assert grouped[0]["referenced_evidence_ids"] == [first_id, second_id]
    assert grouped[0]["tags"] == ["repeat", "threat"]

def test_build_timeline_prototype_processes_message_image_with_injected_ocr():
    image_path = _write_test_file(f"{uuid4()}-message.png", b"fake-image")

    def fake_ocr_runner(_image_path: str) -> OCRResult:
        return OCRResult(
            full_text="2026-03-17 18:30 repeated threatening message",
            segments=[
                OCRSegment(
                    text="2026-03-17 18:30 repeated threatening message",
                )
            ],
            language="ko",
            engine="fake-ocr",
        )

    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="MESSAGE",
                file_format="IMAGE",
                file_name="message.png",
                file_bytes=image_path.read_bytes(),
            ),
        ],
    )

    result = build_timeline_prototype(
        payload,
        llm_client=MockLLMClient(),
        ocr_runner=fake_ocr_runner,
    )

    evidence_result = result.evidence_results[0]
    assert evidence_result.status == "completed"
    assert evidence_result.source_type == "ocr"
    assert evidence_result.normalized_text is not None
    assert result.items[0].date == "2026-03-17"

def test_build_timeline_prototype_combines_date_and_time_from_message_full_text():
    image_path = _write_test_file(f"{uuid4()}-message-multi-line.png", b"fake-image")

    def fake_ocr_runner(_image_path: str) -> OCRResult:
        return OCRResult(
            full_text="2026년 3월 17일 금요일\n오후 11:20\n위협 메시지가 반복적으로 도착했다",
            segments=[
                OCRSegment(text="2026년 3월 17일 금요일"),
                OCRSegment(text="오후 11:20"),
                OCRSegment(text="위협 메시지가 반복적으로 도착했다"),
            ],
            language="ko",
            engine="fake-ocr",
        )

    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="MESSAGE",
                file_format="IMAGE",
                file_name="message.png",
                file_bytes=image_path.read_bytes(),
            ),
        ],
    )

    result = build_timeline_prototype(
        payload,
        llm_client=MockLLMClient(),
        ocr_runner=fake_ocr_runner,
    )

    evidence_result = result.evidence_results[0]
    assert evidence_result.timestamp is not None
    assert evidence_result.timestamp.isoformat() == "2026-03-17T23:20:00"
    assert result.items[0].date == "2026-03-17"
    assert result.items[0].events[0].time == "23:20"

def test_build_timeline_prototype_uses_file_created_at_when_text_has_no_timestamp():
    image_path = _write_test_file(f"{uuid4()}-message-no-timestamp.png", b"fake-image")

    def fake_ocr_runner(_image_path: str) -> OCRResult:
        return OCRResult(
            full_text="위협적인 메시지가 반복적으로 도착했다",
            segments=[
                OCRSegment(text="위협적인 메시지가 반복적으로 도착했다"),
            ],
            language="ko",
            engine="fake-ocr",
        )

    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="MESSAGE",
                file_format="IMAGE",
                file_name="message.png",
                file_bytes=image_path.read_bytes(),
                file_created_at=datetime(2026, 3, 18, 8, 45),
            ),
        ],
    )

    result = build_timeline_prototype(
        payload,
        llm_client=MockLLMClient(),
        ocr_runner=fake_ocr_runner,
    )

    evidence_result = result.evidence_results[0]
    assert evidence_result.timestamp is not None
    assert evidence_result.timestamp.isoformat() == "2026-03-18T08:45:00"
    assert result.items[0].date == "2026-03-18"
    assert result.items[0].events[0].time == "08:45"

def test_build_timeline_prototype_processes_voice_audio_with_injected_stt():
    audio_path = _write_test_file(f"{uuid4()}-voice.m4a", b"fake-audio")

    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="VOICE",
                file_format="AUDIO",
                file_name="voice.m4a",
                file_bytes=audio_path.read_bytes(),
            ),
        ],
    )

    result = build_timeline_prototype(
        payload,
        llm_client=MockLLMClient(),
        stt_engine=MockSTT(),
    )

    evidence_result = result.evidence_results[0]
    assert evidence_result.status == "completed"
    assert evidence_result.source_type == "stt"
    assert evidence_result.normalized_text is not None
    assert evidence_result.normalized_text.endswith("voice.m4a")

def test_build_timeline_prototype_processes_docx_report_record():
    pytest.importorskip("docx")
    from docx import Document

    docx_path = _write_test_file(f"{uuid4()}-record.docx", b"")

    document = Document()
    document.add_paragraph("2026-03-15 consultation record")
    document.add_paragraph("repeated contact was documented")
    document.save(str(docx_path))

    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="REPORT_RECORD",
                file_format="DOCX",
                file_name="record.docx",
                file_bytes=docx_path.read_bytes(),
            ),
        ],
    )

    result = build_timeline_prototype(payload, llm_client=MockLLMClient())

    evidence_result = result.evidence_results[0]
    assert evidence_result.status == "completed"
    assert evidence_result.source_type == "document"
    assert evidence_result.normalized_text is not None
    assert "consultation record" in evidence_result.normalized_text
    assert result.items[0].date == "2026-03-15"

class TagOnlyLLMClient:
    def generate(self, messages: list[dict]) -> str:
        return json.dumps(
            {
                "evidence_metadata": {
                    "value": {
                        "evidence_type": "text",
                        "source": "unknown",
                        "sources": ["unknown"],
                        "created_at": "unknown",
                    },
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "parties": {
                    "value": {
                        "actor": "unknown",
                        "target": "unknown",
                        "relationship": "unknown",
                    },
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "period": {
                    "value": "unknown",
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "frequency": {
                    "value": "unknown",
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "channel": {
                    "value": ["unknown"],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "locations": {
                    "value": ["unknown"],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "action_types": {
                    "value": [],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "refusal_signal": {
                    "value": "unknown",
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "threat_indicators": {
                    "value": [],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "impact_on_victim": {
                    "value": [],
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "report_or_record": {
                    "value": "unknown",
                    "confidence": "low",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "tags": {
                    "value": ["repeat", "threat"],
                    "confidence": "medium",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
                "timeline_summary": {
                    "value": {
                        "title": "tag test title",
                        "description": "tag test description",
                    },
                    "confidence": "medium",
                    "evidence_span": None,
                    "evidence_anchor": None,
                },
            },
            ensure_ascii=False,
        )

def test_build_timeline_prototype_prefers_llm_tags():
    payload = TimelinePrototypeAiInput(
        complaint_id=uuid4(),
        evidences=[
            TimelinePrototypeEvidenceInput(
                evidence_id=uuid4(),
                type="REPORT_RECORD",
                file_format="TXT",
                extracted_text="Repeated threatening contact was documented.",
            ),
        ],
    )

    result = build_timeline_prototype(payload, llm_client=TagOnlyLLMClient())

    evidence_result = result.evidence_results[0]
    timeline_evidence = result.items[0].events[0].evidences[0]

    assert evidence_result.tags == ["repeat", "threat"]
    assert timeline_evidence.tags == ["repeat", "threat"]