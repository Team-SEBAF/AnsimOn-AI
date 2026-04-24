from io import BytesIO
import re
from typing import Optional
from zipfile import ZipFile

from PIL import Image, UnidentifiedImageError

from ansimon_ai.ocr.from_ocr import ocr_image_to_result
from ansimon_ai.ocr.table_formatting import format_ocr_result_text

def _normalize_line(text: str) -> str:
    text = text.replace("\r", " ").replace("\n", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text

def _dedupe_preserve_order(lines: list[str]) -> list[str]:
    normalized_lines: list[str] = []
    seen: set[str] = set()

    for line in lines:
        if not line or line in seen:
            continue
        normalized_lines.append(line)
        seen.add(line)

    return normalized_lines

def _append_normalized_lines(lines: list[str], text: str) -> None:
    for raw_line in text.splitlines():
        normalized = _normalize_line(raw_line)
        if normalized:
            lines.append(normalized)

def _extract_text_from_docx_images(
    docx_path: str,
    *,
    engine: Optional[str] = None,
    lang: str = "kor",
) -> list[str]:
    image_texts: list[str] = []
    supported_suffixes = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}

    with ZipFile(docx_path) as archive:
        image_names = sorted(
            name
            for name in archive.namelist()
            if name.startswith("word/media/")
            and any(name.lower().endswith(suffix) for suffix in supported_suffixes)
        )

        for image_name in image_names:
            image_bytes = archive.read(image_name)
            try:
                with Image.open(BytesIO(image_bytes)) as image:
                    result = ocr_image_to_result(image, engine=engine, lang=lang)
            except (UnidentifiedImageError, OSError):
                continue

            formatted_text = format_ocr_result_text(result)
            if formatted_text.strip():
                _append_normalized_lines(image_texts, formatted_text)

    return image_texts

def extract_text_from_docx(
    docx_path: str,
    *,
    engine: Optional[str] = None,
    lang: str = "kor",
) -> str:
    from docx import Document

    document = Document(docx_path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = _normalize_line(paragraph.text)
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [_normalize_line(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                lines.append(" | ".join(cells))

    lines.extend(
        _extract_text_from_docx_images(
            docx_path,
            engine=engine,
            lang=lang,
        )
    )

    return "\n".join(_dedupe_preserve_order(lines))
